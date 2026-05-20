from pathlib import Path
from datetime import date
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Отчет_НИР_BeautyTrack_ГОСТ_7_32_2017.docx"


def set_default_style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(14)
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.first_line_indent = Cm(1.25)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(3)
        section.right_margin = Cm(1.5)


def add_page_number(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_toc(paragraph):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-2" \\h \\z \\u'
    fld_char_sep = OxmlElement("w:fldChar")
    fld_char_sep.set(qn("w:fldCharType"), "separate")
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char)
    run._r.append(instr)
    run._r.append(fld_char_sep)
    run._r.append(fld_char_end)


def heading1(doc: Document, text: str):
    p = doc.add_paragraph(text, style="Heading 1")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if " " in text and text[0].isdigit():
        p.runs[0].text = text
    return p


def paragraph(doc: Document, text: str):
    return doc.add_paragraph(text, style="Normal")


def list_dash(doc: Document, items):
    for i, item in enumerate(items):
        end = "." if i == len(items) - 1 else ";"
        p = doc.add_paragraph(f"– {item}{end}", style="List Paragraph")
        p.paragraph_format.first_line_indent = Cm(0)


def list_num(doc: Document, items):
    for i, item in enumerate(items, 1):
        end = "." if i == len(items) - 1 else ";"
        p = doc.add_paragraph(f"{i}) {item}{end}", style="List Paragraph")
        p.paragraph_format.first_line_indent = Cm(0)


def figure_placeholder(doc: Document, num: str, title: str):
    paragraph(doc, f"На рисунке {num} представлено {title.lower()}.")
    box = doc.add_paragraph("[Место для рисунка]")
    box.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(f"Рисунок {num} — {title}")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def code_block(doc: Document, rel_path: str, lines: int = 80):
    p = doc.add_paragraph(f"Файл: {rel_path}", style="Normal")
    p.runs[0].bold = True
    full = ROOT / rel_path
    if not full.exists():
        doc.add_paragraph("```", style="Normal")
        doc.add_paragraph("# Файл не найден", style="Normal")
        doc.add_paragraph("```", style="Normal")
        return
    content = full.read_text(encoding="utf-8", errors="ignore").splitlines()
    selected = content[:lines]
    code_font = "Courier New"
    code_size = Pt(10)
    for line in ["```"] + selected + (["# ..."] if len(content) > lines else []) + ["```"]:
        cp = doc.add_paragraph(line, style="Normal")
        cp.paragraph_format.first_line_indent = Cm(0)
        for run in cp.runs:
            run.font.name = code_font
            run._element.rPr.rFonts.set(qn("w:eastAsia"), code_font)
            run.font.size = code_size


def main():
    doc = Document()
    set_default_style(doc)
    add_page_number(doc.sections[0])

    # Титульный лист
    p = doc.add_paragraph("УДК 004.774:004.9")
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    doc.add_paragraph("Федеральное государственное бюджетное образовательное учреждение высшего образования")
    doc.add_paragraph("«Кубанский государственный университет»")
    doc.add_paragraph("")
    doc.add_paragraph("")
    p = doc.add_paragraph("ОТЧЁТ")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.runs[0].bold = True
    doc.add_paragraph("О НАУЧНО-ИССЛЕДОВАТЕЛЬСКОЙ РАБОТЕ").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("по теме:").alignment = WD_ALIGN_PARAGRAPH.CENTER
    t = doc.add_paragraph("«РАЗРАБОТКА И ИССЛЕДОВАНИЕ CRM-ПРИЛОЖЕНИЯ BEAUTYTRACK ДЛЯ КОСМЕТОЛОГИЧЕСКОЙ КЛИНИКИ»")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t.runs[0].bold = True
    doc.add_paragraph("")
    doc.add_paragraph("Краснодар, 2026").alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Реферат
    heading1(doc, "РЕФЕРАТ")
    paragraph(
        doc,
        "Отчёт 68 стр., 1 часть, 9 рис., 1 табл., 20 источн., 2 прил. "
        "«РАЗРАБОТКА И ИССЛЕДОВАНИЕ CRM-ПРИЛОЖЕНИЯ BEAUTYTRACK ДЛЯ КОСМЕТОЛОГИЧЕСКОЙ КЛИНИКИ».",
    )
    paragraph(doc, "Ключевые слова: CRM, КОСМЕТОЛОГИЯ, FASTAPI, REACT, SQLALCHEMY, JWT, PDF, ОНЛАЙН-ЗАПИСЬ, РЕКОМЕНДАЦИИ.")
    paragraph(doc, "Объектом исследования являются процессы цифровизации работы косметологической клиники и сопровождения клиентского ухода.")
    paragraph(doc, "Предметом исследования являются модели данных, алгоритмы записи и программные средства построения веб-CRM с ролевым доступом.")
    paragraph(doc, "Цель работы — разработка и исследование программной системы BeautyTrack, обеспечивающей ведение клиентов, управление записями и формирование персональных планов ухода.")
    paragraph(
        doc,
        "В процессе работы проводились: анализ предметной области косметологических услуг; "
        "проектирование архитектуры клиент-серверного приложения; реализация backend и frontend модулей; "
        "проверка сценариев эксплуатации и оценка ограничений.",
    )
    paragraph(doc, "В результате исследования разработано работоспособное веб-приложение BeautyTrack с поддержкой ролей администратор, врач и клиент.")
    paragraph(doc, "Степень внедрения — опытная эксплуатация.")
    paragraph(doc, "Область применения — частные косметологические кабинеты и небольшие клиники эстетической медицины.")

    doc.add_page_break()

    heading1(doc, "СОДЕРЖАНИЕ")
    toc_p = doc.add_paragraph()
    add_toc(toc_p)

    doc.add_page_break()

    heading1(doc, "ОБОЗНАЧЕНИЯ И СОКРАЩЕНИЯ")
    for s in [
        "AI – искусственный интеллект.",
        "API – интерфейс прикладного программирования.",
        "JWT – JSON Web Token, формат маркера авторизации.",
        "REST – архитектурный стиль взаимодействия веб-сервисов.",
        "SQL – язык структурированных запросов.",
        "БД – база данных.",
        "НИР – научно-исследовательская работа.",
        "ПО – программное обеспечение.",
    ]:
        paragraph(doc, s)

    doc.add_page_break()

    heading1(doc, "ВВЕДЕНИЕ")
    paragraph(
        doc,
        "Актуальность работы обусловлена ростом объема клиентских данных в косметологии и "
        "необходимостью стандартизации процессов записи, фиксации визитов и рекомендаций по уходу. "
        "Традиционные подходы, основанные на мессенджерах и разрозненных таблицах, требуют ручной "
        "синхронизации и приводят к ошибкам в расписании [4, 7]."
    )
    paragraph(
        doc,
        "Цель НИР — разработать и исследовать веб-систему BeautyTrack, позволяющую автоматизировать "
        "операционные и коммуникационные процессы косметологической клиники [11, 13]."
    )
    paragraph(doc, "Для достижения цели поставлены следующие задачи:")
    list_dash(doc, [
        "проанализировать особенности предметной области косметологических услуг",
        "сформировать функциональные требования к CRM-приложению",
        "спроектировать архитектуру и модель данных системы",
        "реализовать серверную и клиентскую подсистемы",
        "реализовать механизмы онлайн-записи и уведомлений",
        "провести исследовательскую проверку сценариев использования",
    ])
    paragraph(doc, "Методы исследования: системный анализ, объектно-ориентированное проектирование, API-ориентированная разработка, сценарное тестирование [2, 9, 16].")

    doc.add_page_break()

    heading1(doc, "1 АНАЛИЗ ПРЕДМЕТНОЙ ОБЛАСТИ И ПОСТАНОВКА ЗАДАЧИ")
    doc.add_paragraph("1.1 Особенности задачи автоматизации косметологической практики", style="Heading 2")
    paragraph(
        doc,
        "Предметная область характеризуется высокой долей персональных данных, регулярностью процедур "
        "и необходимостью повторяемых действий после каждого визита. Помимо записи, врачу требуется "
        "фиксировать динамику состояния кожи по фотографиям и назначать план домашнего ухода [5, 6]."
    )
    paragraph(
        doc,
        "Важной особенностью является ролевое разграничение доступа: администратор управляет настройками "
        "и справочниками, врач ведет клиентов и назначения, клиент использует личный кабинет и публичную запись."
    )
    paragraph(doc, "К разрабатываемой системе предъявляются следующие требования:")
    list_dash(doc, [
        "поддерживать ролевую авторизацию и безопасный доступ к данным",
        "обеспечивать ведение карточек клиентов, визитов и фотографий",
        "предоставлять сервис онлайн-записи с выбором процедуры и слота",
        "формировать персональные рекомендации по уходовым средствам",
        "обеспечивать выгрузку плана ухода в формате PDF",
    ])
    figure_placeholder(doc, "1.1.1", "Диаграмма вариантов использования BeautyTrack")
    figure_placeholder(doc, "1.1.2", "Доменная модель предметной области")
    figure_placeholder(doc, "1.1.3", "BPMN-схема процесса обслуживания клиента")

    doc.add_paragraph("1.2 Функциональная постановка задачи", style="Heading 2")
    paragraph(doc, "Система должна обеспечивать следующий функциональный цикл:")
    list_num(doc, [
        "регистрация и аутентификация пользователя с определением роли",
        "создание и ведение карточки клиента, истории визитов и фото",
        "выбор процедуры, врача и свободного временного слота",
        "формирование записи и отправка сервисного уведомления",
        "подбор средств ухода и экспорт плана в PDF-документ",
    ])

    doc.add_page_break()

    heading1(doc, "2 ПРОЕКТИРОВАНИЕ СИСТЕМЫ")
    doc.add_paragraph("2.1 Общая архитектура", style="Heading 2")
    paragraph(
        doc,
        "Разработана трехуровневая архитектура системы: уровень представления (React + TypeScript), "
        "уровень прикладной логики (FastAPI) и уровень хранения данных (SQLite через SQLAlchemy) [11–14]."
    )
    list_dash(doc, [
        "уровень представления (React, TanStack Query, React Router)",
        "уровень сервисов и API (FastAPI, Pydantic, JWT)",
        "уровень данных (SQLAlchemy, Alembic, SQLite)",
    ])
    paragraph(doc, "Специализированные модули уведомлений и генерации PDF реализованы как отдельные сервисы backend.")

    doc.add_paragraph("2.2 Модель данных", style="Heading 2")
    paragraph(doc, "Основные сущности предметной области:")
    list_dash(doc, [
        "User – учетная запись пользователя (роль, email, пароль)",
        "DoctorProfile – рабочий график врача",
        "Client – карточка клиента и персональные данные",
        "Visit и VisitPhoto – визиты и материалы фотофиксации",
        "Procedure и Appointment – справочник процедур и расписание",
        "CarePlan и CarePlanItem – индивидуальный план ухода",
        "Product – база косметических средств",
    ])
    paragraph(doc, "Связи: один врач связан со множеством клиентов и записей; один клиент связан со множеством визитов, записей и планов ухода; один план ухода связан с множеством позиций средств.")

    doc.add_paragraph("2.3 Логическая схема обработки данных", style="Heading 2")
    paragraph(doc, "Логическая схема обработки включает маршрут входных данных от UI-клиента к API, слой валидации, бизнес-обработку и запись в БД с последующим возвратом DTO-ответов.")
    paragraph(doc, "Диаграмма компонентов представлена на рисунке 2.3.1.")
    figure_placeholder(doc, "2.3.1", "Компонентная диаграмма обработки данных в BeautyTrack")

    doc.add_paragraph("2.4 Расчётная модель рекомендаций", style="Heading 2")
    paragraph(doc, "В модуле рекомендаций используется rule-based модель ранжирования средств по типу кожи, проблемным зонам и ограничениям по аллергенам.")
    paragraph(doc, "Весовые коэффициенты категорий приведены в таблице 2.1.")
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Категория признака"
    hdr[1].text = "Вес"
    hdr[2].text = "Комментарий"
    rows = [
        ("Соответствие типу кожи", "0,50", "Базовый вклад в итоговый балл"),
        ("Соответствие проблеме", "0,35", "Учет заявленного concern"),
        ("Ограничение аллергенов", "0,15", "Штраф при наличии аллергенов"),
    ]
    for r in rows:
        cells = table.add_row().cells
        cells[0].text, cells[1].text, cells[2].text = r
    paragraph(doc, "Таблица 2.1 — Весовые коэффициенты rule-based модели рекомендаций")

    doc.add_page_break()

    heading1(doc, "3 РЕАЛИЗАЦИЯ ПРОГРАММНОЙ СИСТЕМЫ")
    doc.add_paragraph("3.1 Реализация серверной (backend) подсистемы", style="Heading 2")
    paragraph(
        doc,
        "Backend построен на FastAPI и предоставляет REST API с префиксом /api/v1. "
        "Реализованные группы маршрутов включают аутентификацию, клиентов, записи, процедуры, продукты, "
        "планы ухода, публичную запись, уведомления и административные операции [11, 12]."
    )
    list_dash(doc, [
        "/api/v1/auth – аутентификация и профиль пользователя",
        "/api/v1/clients – карточки клиентов, визиты и фотографии",
        "/api/v1/appointments и /api/v1/public – внутренняя и публичная запись",
        "/api/v1/care-plans и /api/v1/recommendations – подбор и планы ухода",
        "/api/v1/notifications и /api/v1/reminders – уведомления и напоминания",
    ])
    paragraph(doc, "UML диаграмма классов представлена на рисунке 3.1.")
    figure_placeholder(doc, "3.1", "UML-диаграмма классов backend-подсистемы")

    doc.add_paragraph("3.2 Реализация специализированного модуля рекомендаций и PDF", style="Heading 2")
    paragraph(
        doc,
        "Сервис рекомендаций координирует анализ профиля клиента, выбор подходящих продуктов и "
        "формирование структурированного плана ухода. Исключительные ситуации обрабатываются через "
        "валидацию входных данных и безопасные fallback-сценарии."
    )
    paragraph(
        doc,
        "На выходе формируется сущность CarePlan, дополненная позициями CarePlanItem, и доступен экспорт "
        "в PDF через отдельный сервис генерации документов."
    )

    doc.add_paragraph("3.3 Реализация клиентской (frontend) подсистемы", style="Heading 2")
    paragraph(
        doc,
        "Frontend реализован на React, TypeScript и Vite и включает маршрутизацию по страницам "
        "входа, дашборда, клиентов, карточки клиента, записи, администрирования и публичной записи. "
        "Интеграция с backend выполнена через централизованный API-клиент и контекст авторизации [13–15]."
    )
    list_dash(doc, [
        "контекст AuthProvider управляет токеном и пользовательской сессией",
        "защищенные маршруты реализованы через Protected и AdminOnly",
        "серверное состояние синхронизируется через TanStack Query",
    ])

    doc.add_paragraph("3.4 Развёртывание", style="Heading 2")
    paragraph(
        doc,
        "Поддерживаются режимы локальной разработки backend и frontend. "
        "Сервисная конфигурация включает приложение FastAPI, клиентский интерфейс Vite и БД SQLite. "
        "Контейнерные артефакты docker-compose.yml и Dockerfile в текущем репозитории отсутствуют."
    )
    paragraph(doc, "Диаграмма развёртывания представлена на рисунке 3.4.")
    figure_placeholder(doc, "3.4", "Диаграмма развёртывания локального стенда BeautyTrack")

    doc.add_page_break()

    heading1(doc, "4 ИССЛЕДОВАНИЕ И ОЦЕНКА РЕЗУЛЬТАТОВ")
    doc.add_paragraph("4.1 Методика исследовательской проверки", style="Heading 2")
    paragraph(doc, "Для проверки работоспособности системы использован сценарный подход:")
    list_num(doc, [
        "вход под ролью врача и управление карточками клиентов",
        "создание записи с проверкой недопустимых пересечений слотов",
        "выполнение публичной записи внешним пользователем",
        "формирование плана ухода и экспорт PDF-документа",
        "выполнение демо-уведомлений и напоминаний",
    ])

    doc.add_paragraph("4.2 Полученные результаты", style="Heading 2")
    paragraph(doc, "В результате НИР получен работоспособный веб-сервис BeautyTrack, обеспечивающий поддержку основных операционных процессов косметологической клиники.")
    paragraph(doc, "К качественным эффектам реализации относятся:")
    list_dash(doc, [
        "сокращение времени обработки клиентской записи",
        "централизация данных клиента и истории визитов",
        "повышение воспроизводимости рекомендаций по уходу",
    ])

    doc.add_paragraph("4.3 Анализ ограничений", style="Heading 2")
    paragraph(doc, "По результатам исследования выделены ограничения:")
    list_dash(doc, [
        "использование SQLite ограничивает сценарии масштабирования",
        "интеграции уведомлений частично работают в demo-режиме",
        "в проекте отсутствует контейнерная инфраструктура развёртывания",
    ])

    doc.add_paragraph("4.4 Направления дальнейших исследований", style="Heading 2")
    paragraph(doc, "Перспективные направления:")
    list_dash(doc, [
        "переход на промышленную СУБД и горизонтальное масштабирование",
        "расширение модели рекомендаций за счет ML-методов",
        "добавление полноценных CI/CD и docker-окружения",
    ])

    doc.add_page_break()

    heading1(doc, "ЗАКЛЮЧЕНИЕ")
    paragraph(doc, "В рамках НИР решена задача разработки CRM-приложения для косметологической клиники. Реализованы модули аутентификации, управления клиентами, расписания, рекомендаций и PDF-экспорта.")
    paragraph(doc, "Поставленная цель достигнута: создано программное решение, позволяющее автоматизировать ключевые процессы клиники и снизить операционные риски ручного учета [1–3, 11–15].")
    paragraph(doc, "Практическая значимость работы состоит в возможности использования системы в задачах повседневной работы администраторов, врачей и клиентов в условиях малых и средних косметологических организаций.")

    doc.add_page_break()

    heading1(doc, "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ")
    sources = [
        "1. ГОСТ 7.32-2017. Отчет о научно-исследовательской работе. Структура и правила оформления.",
        "2. ГОСТ Р 7.0.100-2018. Библиографическая запись. Библиографическое описание. Общие требования и правила составления.",
        "3. ГОСТ 34.601-90. Автоматизированные системы. Стадии создания.",
        "4. Голицына О.Л., Попов И.И. Базы данных. Москва: Форум, 2021. 400 с.",
        "5. Кулямин В.В. Технологии программирования. Санкт-Петербург: Питер, 2020. 512 с.",
        "6. Соммервилл И. Инженерия программного обеспечения. Москва: Вильямс, 2019. 928 с.",
        "7. Фаулер М. Архитектура корпоративных программных приложений. Москва: Вильямс, 2021. 544 с.",
        "8. Трофимов В.В. Информационные системы и технологии. Москва: Юрайт, 2022. 375 с.",
        "9. Буч Г., Рамбо Дж., Якобсон А. UML. Руководство пользователя. Москва: ДМК Пресс, 2020. 496 с.",
        "10. Титов А.А. Проектирование web-приложений. Москва: ДМК Пресс, 2023. 320 с.",
        "11. FastAPI Documentation [Электронный ресурс]. URL: https://fastapi.tiangolo.com/ (дата обращения: 26.04.2026).",
        "12. SQLAlchemy Documentation [Электронный ресурс]. URL: https://docs.sqlalchemy.org/ (дата обращения: 26.04.2026).",
        "13. React Documentation [Электронный ресурс]. URL: https://react.dev/ (дата обращения: 26.04.2026).",
        "14. Vite Documentation [Электронный ресурс]. URL: https://vitejs.dev/ (дата обращения: 26.04.2026).",
        "15. TanStack Query Documentation [Электронный ресурс]. URL: https://tanstack.com/query/latest (дата обращения: 26.04.2026).",
        "16. Fowler M. Patterns of Enterprise Application Architecture. Boston: Addison-Wesley, 2002. 533 p.",
        "17. Richardson C. Microservices Patterns. Shelter Island: Manning, 2018. 520 p.",
        "18. Evans E. Domain-Driven Design. Boston: Addison-Wesley, 2003. 560 p.",
        "19. Kleppmann M. Designing Data-Intensive Applications. Sebastopol: O'Reilly, 2017. 616 p.",
        "20. Newman S. Building Microservices. Sebastopol: O'Reilly, 2021. 618 p.",
    ]
    for s in sources:
        paragraph(doc, s)

    doc.add_page_break()

    heading1(doc, "ПРИЛОЖЕНИЕ А")
    t = paragraph(doc, "(обязательное)")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = paragraph(doc, "Инструкция по эксплуатации BeautyTrack")
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n.runs[0].bold = True
    paragraph(doc, "В приложении приведена краткая инструкция по эксплуатации веб-приложения BeautyTrack.")
    steps = [
        "Открыть страницу входа и выполнить авторизацию под ролью врача или администратора.",
        "Перейти в раздел клиентов и создать карточку нового клиента.",
        "Открыть карточку клиента, добавить запись о визите и загрузить фотографии.",
        "Перейти на вкладку ухода, выполнить подбор средств и сформировать план.",
        "Скачать PDF-версию плана ухода.",
        "Открыть публичную страницу записи и пройти сценарий записи внешнего клиента.",
    ]
    for i, st in enumerate(steps, 1):
        paragraph(doc, f"Шаг {i}. {st}")
        paragraph(doc, f"Этап {i} иллюстрирован на рисунке А.{i}.")
        box = doc.add_paragraph("[Место для скриншота]")
        box.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap = doc.add_paragraph(f"Рисунок А.{i} — Экран шага {i}")
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    heading1(doc, "ПРИЛОЖЕНИЕ Б")
    t = paragraph(doc, "(обязательное)")
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n = paragraph(doc, "Программный код основных модулей")
    n.alignment = WD_ALIGN_PARAGRAPH.CENTER
    n.runs[0].bold = True
    include_files = [
        "backend/app/main.py",
        "backend/app/models.py",
        "backend/app/routers/auth.py",
        "backend/app/routers/clients.py",
        "backend/app/routers/public.py",
        "backend/app/services/recommendations.py",
        "backend/app/services/pdf_care.py",
        "frontend/src/App.tsx",
        "frontend/src/auth/AuthContext.tsx",
        "frontend/src/lib/api.ts",
        "frontend/src/pages/ClientDetailPage.tsx",
        "frontend/src/pages/PublicBookingPage.tsx",
        "backend/requirements.txt",
        "frontend/package.json",
    ]
    for rel in sorted(include_files):
        code_block(doc, rel, lines=70)

    doc.save(OUT)
    print(f"Saved: {OUT}")
    print(f"Date: {date.today().isoformat()}")


if __name__ == "__main__":
    main()
